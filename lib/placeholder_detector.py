"""
Placeholder detection for Legal Document Filler.
Detects multiple placeholder patterns in .docx files.
"""

import logging
import re
from docx import Document
from typing import Dict, List, Set


logger = logging.getLogger(__name__)


class PlaceholderDetectionError(Exception):
    """Custom exception for placeholder detection errors."""
    pass


def detect_placeholders(docx_path: str) -> Dict[str, List[str]]:
    """
    Detect placeholders in a .docx document using multiple patterns.
    
    Patterns supported:
    - {{placeholder_name}} - Double curly braces
    - {placeholder_name} - Single curly braces
    - [Placeholder Name] - Square brackets (any case, e.g., [Date of Safe])
    - _____ - Underscores (3 or more)
    - $[_____] - Dollar sign with brackets and underscores
    
    Args:
        docx_path (str): Path to the .docx file
    
    Returns:
        dict: {normalized_name: [original_patterns]}
              Example: {'client_name': ['{{client_name}}', '{CLIENT_NAME}']}
    
    Raises:
        PlaceholderDetectionError: If document cannot be parsed
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise PlaceholderDetectionError(f"Failed to open document: {str(e)}")
    
    placeholders = {}  # {normalized_name: [original_patterns]}
    
    # Define regex patterns (more permissive; support mixed case and spacing)
    patterns = [
        (r'\$\s*\[\s*_{3,}\s*\]', 'dollar_underscore'),               # $[_____]
        (r'\[\s*([A-Za-z][A-Za-z0-9_\s-]+?)\s*\]', 'square_bracket'),  # [Placeholder Name]
        (r'\{\{\s*([A-Za-z0-9_\s-]+?)\s*\}\}', 'double_curly'),      # {{ placeholder }}
        (r'\{\s*([A-Za-z0-9_\s-]+?)\s*\}', 'single_curly'),            # { placeholder }
        (r'_{3,}', 'underscore'),                                          # ___ (3+)
    ]
    
    # Track context for underscore patterns
    underscore_counter = 0
    
    # Helper to process a single text chunk
    def process_text(text: str):
        nonlocal underscore_counter
        if not text or not text.strip():
            return
        
        # Check each pattern
        for pattern, pattern_type in patterns:
            matches = re.finditer(pattern, text)
            
            for match in matches:
                original = match.group(0)
                
                # Normalize placeholder name
                if pattern_type == 'underscore':
                    # Try to extract label before underscores
                    # Look for pattern like "Name: _____" or "Name:_____"
                    before_match = text[:match.start()].strip()
                    label_match = re.search(r'([A-Za-z][A-Za-z0-9_\s-]{1,50})\s*:\s*$', before_match)
                    
                    if label_match:
                        normalized = normalize_placeholder_name(label_match.group(1))
                    else:
                        # Use generic name with counter
                        underscore_counter += 1
                        normalized = f"field_{underscore_counter}"
                
                elif pattern_type == 'dollar_underscore':
                    # Try to find a descriptive label AFTER the brackets, e.g., (the "Purchase Amount")
                    after_text = text[match.end():match.end()+100]
                    after_label = re.search(r"\(\s*the\s+['\"""]?([A-Za-z][A-Za-z0-9_\s-]{2,})['\"""]?\s*\)", after_text, re.IGNORECASE)

                    if after_label:
                        normalized = normalize_placeholder_name(after_label.group(1))
                    else:
                        # Otherwise, try label immediately before the brackets
                        before_match = text[:match.start()].strip()
                        label_match = re.search(r'([A-Za-z][A-Za-z0-9_\s-]{2,50})\s*:?\s*$', before_match)
                        stopwords = {'of', 'the', 'a', 'an', 'is', 'at', 'on', 'by', 'for', 'in'}
                        if label_match and label_match.group(1).strip().lower() not in stopwords:
                            normalized = normalize_placeholder_name(label_match.group(1))
                        else:
                            base = 'amount'
                            normalized = f"{base}_{len([k for k in placeholders.keys() if k.startswith(base + '_')]) + 1}"
                
                else:
                    # For other patterns, normalize the captured content
                    captured = match.group(1)
                    normalized = normalize_placeholder_name(captured)
                
                # Add to placeholders dict
                if normalized not in placeholders:
                    placeholders[normalized] = []
                
                # Only add if not already in list (avoid duplicates in same paragraph)
                if original not in placeholders[normalized]:
                    placeholders[normalized].append(original)

    # Parse document paragraphs
    for paragraph in doc.paragraphs:
        process_text(paragraph.text)

    # Parse tables (signature blocks and forms are often inside tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_text(paragraph.text)

    # Heuristic: capture common signature labels without explicit placeholders
    # e.g., a paragraph that is exactly "Address:" or "Email:" etc.
    signature_labels = {
        'address:': 'address',
        'email:': 'email',
        'phone:': 'phone',
        'by:': 'by',
    }
    def add_label_placeholder(label_key):
        name = signature_labels[label_key]
        key = normalize_placeholder_name(name)
        placeholders.setdefault(key, []).append(label_key)

    # Re-scan document (paragraphs and table cells) for bare labels
    def scan_for_bare_labels(text: str):
        if not text:
            return
        t = text.strip()
        low = t.lower()
        if low in signature_labels:
            add_label_placeholder(low)

    for paragraph in doc.paragraphs:
        scan_for_bare_labels(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    scan_for_bare_labels(paragraph.text)

    return placeholders


def reduce_false_positives(placeholders: Dict[str, List[str]], doc_text: str = None) -> Dict[str, List[str]]:
    """
    Filter out false positives from detected placeholders.
    
    Filters:
    - Numeric citations like [1], [2], [10]
    - Section references like [Section 2(b)], [2(a)]
    - Single-character brackets like [a], [x]
    - "section" keyword in placeholder names
    - Placeholders that appear only once and are very short (unless multi-word)
    
    Args:
        placeholders (dict): Result from detect_placeholders()
        doc_text (str): Full document text (optional, for additional context)
    
    Returns:
        dict: Filtered placeholders
    """
    filtered = {}
    # Ignore lists and patterns
    ignore_acronyms = {
        'llc', 'usa', 'inc', 'ltd', 'co', 'corp', 'aka', 'dba', 'llp', 'pllc'
    }
    # Allowlist of short, legitimate placeholder names that often appear once
    allow_short_singletons = {
        'by', 'name', 'title', 'email', 'address', 'phone'
    }
    
    for normalized, originals in placeholders.items():
        # Skip if normalized name is only digits
        if normalized.isdigit():
            continue
        
        # Skip if contains "section" keyword (likely a section reference)
        if 'section' in normalized.lower():
            continue
        
        # Skip if normalized name is a single character
        if len(normalized) == 1:
            continue
        
        # Check original patterns for false positive indicators
        has_false_positive = False
        
        for orig in originals:
            # Check if it's a square bracket pattern
            if orig.startswith('[') and orig.endswith(']'):
                content = orig[1:-1].strip()
                
                # Filter numeric citations: [1], [12], etc.
                if content.isdigit():
                    has_false_positive = True
                    break
                
                # Filter section references: [2(a)], [Section 2(b)], [Section 3]
                if re.match(r'^\d+\([a-z]\)$', content, re.IGNORECASE):
                    has_false_positive = True
                    break
                
                if re.match(r'^Section\s+\d+', content, re.IGNORECASE):
                    has_false_positive = True
                    break

                # Filter exhibit/schedule/annex tags: [Exhibit A], [Schedule 1], [Annex B]
                if re.match(r'^(Exhibit|Schedule|Annex)\s+[A-Za-z0-9]+$', content, re.IGNORECASE):
                    has_false_positive = True
                    break

                # Filter common acronyms specifically called out (e.g., [LLC], [USA])
                if content.strip().lower() in ignore_acronyms:
                    has_false_positive = True
                    break

                # Filter common email-template boilerplate markers inside brackets
                if re.search(r'unsubscribe|manage\s+preferences|view\s+in\s+browser', content, re.IGNORECASE):
                    has_false_positive = True
                    break
                
                # Filter single character brackets: [a], [x]
                if len(content) == 1:
                    has_false_positive = True
                    break
        
        if has_false_positive:
            continue
        
        # Apply occurrence-based filtering
        occurrence_count = len(originals)
        
        # If appears only once, be more strict
        if occurrence_count == 1:
            # Keep if it's multi-word (has underscore in normalized name)
            if '_' not in normalized:
                # Skip single-occurrence, single-word placeholders that are very short
                if len(normalized) <= 3 and normalized not in allow_short_singletons:
                    continue
        
        # If appears 2+ times or is multi-word, keep it
        if occurrence_count >= 2 or '_' in normalized or len(normalized) > 3:
            filtered[normalized] = originals
    
    return filtered


def get_placeholder_count(placeholders: Dict[str, List[str]]) -> Dict[str, int]:
    """
    Get count of occurrences for each placeholder.
    
    Args:
        placeholders (dict): Result from detect_placeholders()
    
    Returns:
        dict: {normalized_name: count}
    """
    return {name: len(originals) for name, originals in placeholders.items()}


def get_total_occurrences(placeholders: Dict[str, List[str]]) -> int:
    """
    Get total number of placeholder occurrences.
    
    Args:
        placeholders (dict): Result from detect_placeholders()
    
    Returns:
        int: Total occurrence count
    """
    return sum(len(originals) for originals in placeholders.values())


def normalize_placeholder_name(name: str) -> str:
    """
    Normalize a placeholder name to a consistent format.
    
    Rules:
    - Convert to lowercase
    - Replace spaces with underscores
    - Remove special characters except underscores
    
    Args:
        name (str): Original placeholder name
    
    Returns:
        str: Normalized name
    """
    # Convert to lowercase
    normalized = name.lower()
    
    # Replace spaces with underscores
    normalized = normalized.replace(' ', '_')
    
    # Remove special characters except underscores and alphanumeric
    normalized = re.sub(r'[^a-z0-9_]', '', normalized)
    
    # Replace multiple underscores with single
    normalized = re.sub(r'_+', '_', normalized)
    
    # Remove leading/trailing underscores
    normalized = normalized.strip('_')
    
    return normalized


def group_similar_placeholders(placeholders: Dict[str, List[str]]) -> Dict[str, Dict[str, any]]:
    """
    Group placeholders with metadata.
    
    Args:
        placeholders (dict): Result from detect_placeholders()
    
    Returns:
        dict: Enhanced placeholder info with metadata
              {
                  'client_name': {
                      'normalized': 'client_name',
                      'variants': ['{{client_name}}', '{CLIENT_NAME}'],
                      'count': 2,
                      'patterns': ['double_curly', 'single_curly']
                  }
              }
    """
    grouped = {}
    
    for normalized, originals in placeholders.items():
        # Determine pattern types
        pattern_types = set()
        for orig in originals:
            if orig.startswith('{{') and orig.endswith('}}'):
                pattern_types.add('double_curly')
            elif orig.startswith('{') and orig.endswith('}'):
                pattern_types.add('single_curly')
            elif orig.startswith('[') and orig.endswith(']'):
                pattern_types.add('square_bracket')
            elif orig.startswith('$[') or orig.startswith('$ ['):
                pattern_types.add('dollar_underscore')
            elif '_' in orig and not any(c.isalnum() for c in orig):
                pattern_types.add('underscore')
        
        grouped[normalized] = {
            'normalized': normalized,
            'variants': originals,
            'count': len(originals),
            'patterns': list(pattern_types)
        }
    
    return grouped


def get_placeholder_summary(placeholders: Dict[str, List[str]]) -> Dict[str, any]:
    """
    Get a summary of detected placeholders.
    
    Args:
        placeholders (dict): Result from detect_placeholders()
    
    Returns:
        dict: Summary with counts and statistics
    """
    total_unique = len(placeholders)
    total_occurrences = get_total_occurrences(placeholders)
    counts = get_placeholder_count(placeholders)
    grouped = group_similar_placeholders(placeholders)
    
    return {
        'total_unique': total_unique,
        'total_occurrences': total_occurrences,
        'placeholders': list(placeholders.keys()),
        'counts': counts,
        'grouped': grouped
    }


def _split_sentences_with_spans(text: str) -> List[Dict[str, int]]:
    """
    Split text into sentences and return list of dicts with start/end indexes and sentence text.
    Very lightweight heuristic using punctuation . ! ? ; and line breaks.
    """
    if not text:
        return []
    spans = []
    start = 0
    i = 0
    n = len(text)
    terminators = {'.', '!', '?', ';', '\n'}
    while i < n:
        ch = text[i]
        if ch in terminators:
            # Include the terminator
            end = i + 1
            sent = text[start:end].strip()
            if sent:
                spans.append({'start': start, 'end': end, 'text': sent})
            start = end
        i += 1
    # Tail
    if start < n:
        sent = text[start:n].strip()
        if sent:
            spans.append({'start': start, 'end': n, 'text': sent})
    return spans


def _extract_sentence_context(paragraph_text: str, match_start: int, match_end: int) -> Dict[str, str]:
    """
    Extract the sentence containing the match and its immediate previous and next sentences.
    Returns dict with keys: prev, sentence, next.
    """
    sentences = _split_sentences_with_spans(paragraph_text or '')
    if not sentences:
        t = (paragraph_text or '').strip()
        return {'prev': '', 'sentence': t, 'next': ''}
    idx = 0
    for j, span in enumerate(sentences):
        if span['start'] <= match_start < span['end']:
            idx = j
            break
    prev_text = sentences[idx - 1]['text'] if idx - 1 >= 0 else ''
    sent_text = sentences[idx]['text']
    next_text = sentences[idx + 1]['text'] if idx + 1 < len(sentences) else ''
    # Clip to reasonable lengths to keep LLM payload minimal
    def clip(s: str, max_len: int = 300) -> str:
        return s[:max_len]
    return {
        'prev': clip(prev_text),
        'sentence': clip(sent_text, 400),
        'next': clip(next_text)
    }


def detect_placeholders_with_context(docx_path: str) -> Dict[str, any]:
    """
    Detect placeholders and emit candidate metadata with minimal context.
    Uses collect → arbitrate → process pipeline with explicit priorities.
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise PlaceholderDetectionError(f"Failed to open document: {str(e)}")

    placeholders: Dict[str, List[str]] = {}
    candidates: List[Dict[str, any]] = []

    # Patterns and priorities
    patterns = [
        (r'\$\s*\[\s*_{3,}\s*\]', 'dollar_underscore'),
        (r'\[\s*([A-Za-z][A-Za-z0-9_\s-]+?)\s*\]', 'square_bracket'),
        (r'\{\{\s*([A-Za-z0-9_\s-]+?)\s*\}\}', 'double_curly'),
        (r'\{\s*([A-Za-z0-9_\s-]+?)\s*\}', 'single_curly'),
        (r'_{3,}', 'underscore'),
    ]
    priority_map = {
        'signature_label': 6,
        'dollar_underscore': 5,
        'square_bracket': 4,
        'double_curly': 3,
        'single_curly': 2,
        'underscore': 1,
    }

    underscore_counter = 0

    def collect_candidates(text: str, locator: str) -> list:
        collected = []
        if not text:
            return collected
        t = text
        # Pattern-based
        for pattern, pattern_type in patterns:
            for m in re.finditer(pattern, t):
                start, end = m.start(), m.end()
                length = end - start
                original = m.group(0)
                captured = None
                if pattern_type in ('double_curly', 'single_curly', 'square_bracket') and m.groups():
                    captured = m.group(1)
                collected.append({
                    'start': start,
                    'end': end,
                    'length': length,
                    'kind': pattern_type,
                    'priority': priority_map[pattern_type],
                    'original': original,
                    'captured': captured,
                    'locator': locator,
                    'text': t,
                })
        # Signature-line heuristic
        sig_re = re.compile(r'^\s*(Address|Email|E-mail|Phone|Name|Title)\s*:?\s*([ \t\._\-—]*)$', re.IGNORECASE)
        m = sig_re.match(t or '')
        if m:
            remainder = m.group(2) or ''
            if not re.search(r'[A-Za-z0-9]', remainder):
                collected.append({
                    'start': 0,
                    'end': len(t),
                    'length': len(t),
                    'kind': 'signature_label',
                    'priority': priority_map['signature_label'],
                    'original': t.strip(),
                    'captured': m.group(1),
                    'locator': locator,
                    'text': t,
                })
        return collected

    def arbitrate(collected: list) -> list:
        collected.sort(key=lambda c: (c['start'], -c['priority'], -c['length']))
        kept = []
        taken = []  # list of (s,e)
        for c in collected:
            s, e = c['start'], c['end']
            overlap = any(not (e <= s2 or s >= e2) for s2, e2 in taken)
            if overlap:
                logger.debug("[placeholder-detect] dropped overlap kind=%s span=(%d,%d)", c['kind'], s, e)
                continue
            kept.append(c)
            taken.append((s, e))
        return kept

    def process_kept(kept: list):
        nonlocal underscore_counter
        for c in kept:
            t = c['text']
            start, end = c['start'], c['end']
            pattern_type = c['kind']
            original = c['original']

            # Normalize
            if pattern_type == 'underscore':
                before = (t[:start] or '').strip()
                label_m = re.search(r'([A-Za-z][A-Za-z0-9_\s-]{1,50})\s*:\s*$', before)
                if label_m:
                    normalized = normalize_placeholder_name(label_m.group(1))
                else:
                    underscore_counter += 1
                    normalized = f"field_{underscore_counter}"
            elif pattern_type == 'dollar_underscore':
                after_text = t[end:end+100]
                after_label = re.search(r"\(\s*the\s+['\"“”]?([A-Za-z][A-Za-z0-9_\s-]{2,})['\"“”]?\s*\)", after_text, re.IGNORECASE)
                if after_label:
                    normalized = normalize_placeholder_name(after_label.group(1))
                else:
                    before = (t[:start] or '').strip()
                    label_m = re.search(r'([A-Za-z][A-Za-z0-9_\s-]{2,50})\s*:?\s*$', before)
                    stopwords = {'of', 'the', 'a', 'an', 'is', 'at', 'on', 'by', 'for', 'in'}
                    if label_m and label_m.group(1).strip().lower() not in stopwords:
                        normalized = normalize_placeholder_name(label_m.group(1))
                    else:
                        base = 'amount'
                        normalized = f"{base}_{len([k for k in placeholders.keys() if k.startswith(base + '_')]) + 1}"
            elif pattern_type in ('double_curly', 'single_curly', 'square_bracket'):
                normalized = normalize_placeholder_name(c.get('captured') or '')
            elif pattern_type == 'signature_label':
                normalized = normalize_placeholder_name(c.get('captured') or 'field')
            else:
                underscore_counter += 1
                normalized = f"field_{underscore_counter}"

            placeholders.setdefault(normalized, [])
            if original not in placeholders[normalized]:
                placeholders[normalized].append(original)

            context = _extract_sentence_context(t, start, end)
            instance_id = f"{c['locator']}-s{start}-e{end}"
            # Log kept
            def _tr(val: str, lim: int = 120) -> str:
                if not val:
                    return ''
                return val if len(val) <= lim else val[:lim] + '…'
            logger.debug(
                "[placeholder-detect] pattern=%s normalized=%s locator=%s original=%r prev=%r sentence=%r next=%r",
                pattern_type, normalized, c['locator'], original,
                _tr(context.get('prev')), _tr(context.get('sentence')), _tr(context.get('next'))
            )
            candidates.append({
                'id': instance_id,
                'normalized': normalized,
                'original': original,
                'pattern_type': pattern_type,
                'context': context
            })

    # Body paragraphs
    for pi, paragraph in enumerate(doc.paragraphs):
        locator = f"p{pi}"
        kept = arbitrate(collect_candidates(paragraph.text, locator))
        process_kept(kept)

    # Tables
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, paragraph in enumerate(cell.paragraphs):
                    locator = f"t{ti}-r{ri}-c{ci}-p{pi}"
                    kept = arbitrate(collect_candidates(paragraph.text, locator))
                    process_kept(kept)

    groups: Dict[str, List[str]] = {}
    for cand in candidates:
        groups.setdefault(cand['normalized'], []).append(cand['id'])

    return {
        'placeholders': placeholders,
        'candidates': candidates,
        'groups': groups
    }

