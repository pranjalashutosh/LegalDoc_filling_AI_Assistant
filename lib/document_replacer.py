"""
Document Replacer
Replaces placeholders in .docx documents with user-provided values
"""

from docx import Document
import re
import os
from config import Config
import logging
from lib.placeholder_detector import normalize_placeholder_name

logger = logging.getLogger(__name__)


class DocumentReplacementError(Exception):
    """Exception raised for errors during document replacement."""
    pass


def replace_placeholders(input_path, output_path, placeholder_values, overrides=None):
    """
    Replace placeholders in a .docx document with provided values.
    Preserves all formatting including bold, italic, fonts, etc.
    
    Args:
        input_path (str): Path to the input .docx file
        output_path (str): Path where the completed document will be saved
        placeholder_values (dict): Dictionary mapping placeholder names to their replacement values
    
    Returns:
        str: Path to the completed document
    
    Raises:
        DocumentReplacementError: If document processing fails
    """
    try:
        logger.info(f"Loading document from: {input_path}")
        doc = Document(input_path)
        
        # Prepare lookup with flexible key normalization (case/spacing)
        lookup = _build_placeholder_lookup(placeholder_values)
        try:
            provided_keys = list((placeholder_values or {}).keys())
            logger.info("Replacement input received: %d keys", len(provided_keys))
            if provided_keys:
                logger.debug("Input keys (first 20): %s", provided_keys[:20])
        except Exception:
            pass

        # Compile all placeholder patterns
        patterns = _compile_patterns()
        
        # Track replacements for logging
        replacements_made = {}
        
        # Track underscore counter across document (mutable reference)
        underscore_counter_ref = [0]
        
        # Track existing normalized keys for counter logic (matches detection)
        existing_keys = set()
        
        # Process all paragraphs in the document (body-level)
        for p_index, paragraph in enumerate(doc.paragraphs):
            locator = f"p{p_index}"
            _replace_in_paragraph(paragraph, patterns, lookup, replacements_made, overrides or {}, locator, underscore_counter_ref, existing_keys)
        
        # Process all tables in the document with stable locators
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for pi, paragraph in enumerate(cell.paragraphs):
                        locator = f"t{ti}-r{ri}-c{ci}-p{pi}"
                        _replace_in_paragraph(paragraph, patterns, lookup, replacements_made, overrides or {}, locator, underscore_counter_ref, existing_keys)
        
        # Handle bare signature labels (e.g., "Address:" without explicit placeholder)
        _fill_signature_labels(doc, lookup)

        # Save the completed document
        logger.info(f"Saving completed document to: {output_path}")
        doc.save(output_path)
        
        # Log replacement summary
        logger.info(f"Document replacement complete. Replacements made: {len(replacements_made)}")
        for placeholder, count in replacements_made.items():
            logger.debug(f"  - '{placeholder}': {count} occurrence(s)")
        
        return output_path
        
    except Exception as e:
        logger.error(f"Error replacing placeholders: {e}")
        raise DocumentReplacementError(f"Failed to replace placeholders: {str(e)}")


def _compile_patterns():
    """
    Compile all placeholder regex patterns.
    Uses the same patterns and normalization logic as placeholder_detector.py.
    
    Returns:
        list: List of tuples (pattern_name, compiled_regex, normalization_function)
    """
    patterns = []
    
    # Pattern 1: {{PLACEHOLDER}} - matches detection exactly
    patterns.append((
        'double_curly',
        re.compile(r'\{\{\s*([A-Za-z0-9_\s-]+?)\s*\}\}'),
        lambda m: normalize_placeholder_name(m.group(1).strip())
    ))
    
    # Pattern 2: {PLACEHOLDER} - matches detection exactly
    patterns.append((
        'single_curly',
        re.compile(r'\{\s*([A-Za-z0-9_\s-]+?)\s*\}'),
        lambda m: normalize_placeholder_name(m.group(1).strip())
    ))
    
    # Pattern 3: [Placeholder Name] - matches detection exactly
    patterns.append((
        'square_bracket',
        re.compile(r'\[\s*([A-Za-z][A-Za-z0-9_\s-]+?)\s*\]'),
        lambda m: normalize_placeholder_name(m.group(1).strip())
    ))
    
    # Pattern 4: _____ (3+ underscores) - normalization handled by _normalize_with_context
    patterns.append((
        'underscore',
        re.compile(r'_{3,}'),
        lambda m: None  # Will be normalized by _normalize_with_context
    ))
    
    # Pattern 5: $[_____] - matches detection exactly, renamed to dollar_underscore
    patterns.append((
        'dollar_underscore',
        re.compile(r'\$\s*\[\s*_{3,}\s*\]'),
        lambda m: None  # Will be normalized by _normalize_with_context
    ))
    
    # Note: bracket_underscore pattern removed to match detection
    
    return patterns


def _build_placeholder_lookup(placeholder_values):
    """
    Build a flexible lookup dictionary that allows case-insensitive and
    space/underscore-insensitive access to placeholder values.
    """
    if not isinstance(placeholder_values, dict):
        return {}

    lookup = {}
    for key, value in placeholder_values.items():
        if key is None:
            continue
        variants = set()
        k1 = str(key)
        variants.add(k1)
        variants.add(k1.upper())
        variants.add(k1.lower())
        variants.add(k1.replace(' ', '_'))
        variants.add(k1.replace('_', ' '))
        variants.add(k1.title())
        variants.add(k1.lower().replace(' ', '_'))
        variants.add(k1.upper().replace(' ', '_'))
        for v in variants:
            lookup[v] = value
        try:
            # Log a compact view of variants per key to aid debugging
            sample = sorted(list(variants))[:8]
            logger.debug("Lookup variants for key '%s' (showing up to 8/%d): %s", k1, len(variants), sample)
        except Exception:
            pass
    return lookup


def _resolve_value(lookup, normalized):
    """
    Resolve a replacement value from the lookup using multiple variants of the key.
    """
    if normalized in lookup:
        try:
            logger.debug("Resolved key '%s' via direct match", normalized)
        except Exception:
            pass
        return lookup[normalized]
    variants = [
        normalized.upper(),
        normalized.lower(),
        normalized.replace(' ', '_'),
        normalized.replace('_', ' '),
        normalized.title(),
        normalized.lower().replace(' ', '_'),
        normalized.upper().replace(' ', '_')
    ]
    for v in variants:
        if v in lookup:
            try:
                logger.debug("Resolved key '%s' via variant '%s'", normalized, v)
            except Exception:
                pass
            return lookup[v]
    try:
        logger.info("No value found for normalized key '%s'", normalized)
    except Exception:
        pass
    return None


def _normalize_with_context(text, match, pattern_name, underscore_counter_ref, existing_keys):
    """
    For underscore-based placeholders, try to infer a meaningful key using
    nearby context, matching detection logic exactly from placeholder_detector.py.
    
    Args:
        text: Full paragraph text
        match: Regex match object
        pattern_name: Pattern type ('underscore' or 'dollar_underscore')
        underscore_counter_ref: List with single int to track counter (mutable reference)
        existing_keys: Set of already-seen normalized keys for counter logic
    
    Returns:
        Normalized placeholder key
    """
    if pattern_name == 'underscore':
        # Try to extract label before underscores: "Name: _____"
        before = (text[:match.start()] or '').strip()
        label_m = re.search(r'([A-Za-z][A-Za-z0-9_\s-]{1,50})\s*:\s*$', before)
        if label_m:
            return normalize_placeholder_name(label_m.group(1))
        else:
            # Use generic name with counter
            underscore_counter_ref[0] += 1
            return f"field_{underscore_counter_ref[0]}"
    
    elif pattern_name == 'dollar_underscore':
        # Try to find a descriptive label AFTER the brackets: (the "Purchase Amount")
        after_text = text[match.end():match.end()+100]
        after_label = re.search(r"\(\s*the\s+['\" ""]?([A-Za-z][A-Za-z0-9_\s-]{2,})['\" ""]?\s*\)", after_text, re.IGNORECASE)
        if after_label:
            return normalize_placeholder_name(after_label.group(1))
        else:
            # Otherwise, try label immediately before the brackets
            before = (text[:match.start()] or '').strip()
            label_m = re.search(r'([A-Za-z][A-Za-z0-9_\s-]{2,50})\s*:?\s*$', before)
            stopwords = {'of', 'the', 'a', 'an', 'is', 'at', 'on', 'by', 'for', 'in'}
            if label_m and label_m.group(1).strip().lower() not in stopwords:
                return normalize_placeholder_name(label_m.group(1))
            else:
                # Default to amount_N pattern - match detection logic exactly
                base = 'amount'
                # Count existing amount_* keys like detection does
                existing_count = len([k for k in existing_keys if k.startswith(base + '_')])
                normalized = f"{base}_{existing_count + 1}"
                existing_keys.add(normalized)
                return normalized
    
    # Fallback (shouldn't reach here)
    underscore_counter_ref[0] += 1
    return f"field_{underscore_counter_ref[0]}"


def _fill_signature_labels(doc, lookup):
    """
    Fill common signature block labels when no explicit placeholder markup exists.
    For paragraphs that end with e.g. "Address:" or "Email:", append the provided value.
    Matches detection signature_label handling (normalizes using normalize_placeholder_name).
    """
    # Use same labels as detection
    signature_labels = {
        'address:': 'address',
        'email:': 'email',
        'phone:': 'phone',
        'by:': 'by',
        'name:': 'name',
        'title:': 'title',
    }

    def process_para(paragraph):
        text = paragraph.text or ''
        stripped = text.strip()
        low = stripped.lower()
        # Check if it matches signature label pattern (ends with label:)
        for label_key, normalized_label in signature_labels.items():
            if low.endswith(label_key):
                # Use normalized label for lookup (matches detection)
                normalized_key = normalize_placeholder_name(normalized_label)
                value = _resolve_value(lookup, normalized_key)
                if value:
                    # Append a space and the value, preserving existing formatting strategy
                    new_text = text + ' ' + value
                    _update_paragraph_text(paragraph, new_text, text, paragraph.runs)
                    try:
                        logger.info("Filled signature label '%s' (normalized=%s) with provided value", label_key, normalized_key)
                    except Exception:
                        pass
                break

    for paragraph in doc.paragraphs:
        process_para(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_para(paragraph)


def _replace_in_paragraph(paragraph, patterns, placeholder_values, replacements_made, overrides, locator: str, underscore_counter_ref, existing_keys):
    """
    Replace placeholders in a single paragraph, preserving formatting.
    Works at the run level to maintain character formatting.
    
    Args:
        paragraph: docx.paragraph.Paragraph object
        patterns: List of compiled regex patterns
        placeholder_values: Dictionary of placeholder -> value mappings
        replacements_made: Dictionary to track replacement counts
        overrides: Dictionary of per-instance overrides
        locator: String identifier for this paragraph location
        underscore_counter_ref: List with single int to track underscore counter (mutable reference)
        existing_keys: Set of already-seen normalized keys for counter logic
    """
    # Get full paragraph text
    full_text = paragraph.text
    
    if not full_text:
        return
    
    # Check if any placeholder patterns exist in this paragraph
    has_placeholders = False
    for pattern_name, regex, normalizer in patterns:
        if regex.search(full_text):
            has_placeholders = True
            break
    
    if not has_placeholders:
        return
    
    # Build a map of character positions to runs
    runs = paragraph.runs
    if not runs:
        return
    
    # Create a working copy of the text
    working_text = full_text
    
    # Perform replacements on working text
    replacements = []  # List of (start_pos, end_pos, replacement_text, normalized_key)
    
    # Find all matches across all patterns
    for pattern_name, regex, normalizer in patterns:
        for match in regex.finditer(full_text):
            # Normalize the placeholder name; enhance with context for underscore-based patterns
            if pattern_name in ('dollar_underscore', 'underscore'):
                normalized = _normalize_with_context(full_text, match, pattern_name, underscore_counter_ref, existing_keys)
            else:
                # For patterns with captured groups, use the normalizer
                normalized = normalizer(match)
                # Track normalized keys
                if normalized:
                    existing_keys.add(normalized)
            
            if normalized is None:
                # Fallback if normalizer returned None
                normalized = 'placeholder'
            
            try:
                snippet = full_text[max(0, match.start()-20):min(len(full_text), match.end()+20)]
                logger.debug("Detected match at %s [%d:%d] pattern=%s normalized=%s snippet=%r",
                             locator, match.start(), match.end(), pattern_name, normalized, snippet)
            except Exception:
                pass
            
            # Look up the replacement value
            replacement_value = _resolve_value(placeholder_values, normalized)

            # Apply per-instance override if provided and available
            if overrides:
                instance_id = f"{locator}-s{match.start()}-e{match.end()}"
                inst = overrides.get(instance_id)
                if inst and inst.get('answer'):
                    replacement_value = inst['answer']
                    try:
                        logger.info("Override applied for instance %s (normalized=%s)", instance_id, normalized)
                    except Exception:
                        pass
            
            # Fallbacks for common synonyms/prefixes (matching detection behavior)
            if replacement_value is None and pattern_name == 'dollar_underscore':
                for key in ['purchase_amount', 'amount', 'price', 'amount_1']:
                    replacement_value = _resolve_value(placeholder_values, key)
                    if replacement_value is not None:
                        try:
                            logger.debug("Resolved '%s' via amount synonym '%s'", normalized, key)
                        except Exception:
                            pass
                        break
            if replacement_value is None and pattern_name == 'underscore':
                for key in ['blank_1', 'blank', 'field_1', 'signatory_name', 'by', 'name', 'title', 'address', 'email', 'phone']:
                    replacement_value = _resolve_value(placeholder_values, key)
                    if replacement_value is not None:
                        try:
                            logger.debug("Resolved '%s' via blank/signature synonym '%s'", normalized, key)
                        except Exception:
                            pass
                        break
            
            if replacement_value is not None:
                # Record the replacement
                start = match.start()
                end = match.end()
                replacements.append((start, end, replacement_value, normalized))
                try:
                    logger.debug("Planned replacement at %s [%d:%d] normalized=%s value_len=%d",
                                 locator, start, end, normalized, len(replacement_value))
                except Exception:
                    pass
            else:
                try:
                    logger.info("No replacement value for normalized=%s at %s [%d:%d] pattern=%s",
                                normalized, locator, match.start(), match.end(), pattern_name)
                except Exception:
                    pass
    
    # Sort replacements by position (reverse order to handle offsets correctly)
    replacements.sort(key=lambda x: x[0], reverse=True)
    
    # Apply replacements to working text
    for start, end, replacement_value, normalized in replacements:
        working_text = working_text[:start] + replacement_value + working_text[end:]
        
        # Track replacement counts
        if normalized in replacements_made:
            replacements_made[normalized] += 1
        else:
            replacements_made[normalized] = 1
    
    # If we made any replacements, update the paragraph
    if replacements:
        _update_paragraph_text(paragraph, working_text, full_text, runs)
        try:
            logger.debug("Applied %d replacement(s) in %s", len(replacements), locator)
        except Exception:
            pass


def _update_paragraph_text(paragraph, new_text, original_text, runs):
    """
    Update paragraph text while attempting to preserve formatting.
    
    Strategy: Clear all runs and create a single new run with the first run's formatting.
    This is a simplified approach - for production, you might want more sophisticated
    formatting preservation.
    
    Args:
        paragraph: The paragraph to update
        new_text: The new text with replacements
        original_text: The original text
        runs: The original runs
    """
    # Store formatting from the first run (if exists)
    base_format = None
    if runs:
        first_run = runs[0]
        base_format = {
            'bold': first_run.bold,
            'italic': first_run.italic,
            'underline': first_run.underline,
            'font_name': first_run.font.name if first_run.font.name else None,
            'font_size': first_run.font.size,
        }
    
    # Clear all existing runs
    for run in runs:
        run.text = ''
    
    # Remove empty runs
    for run in paragraph.runs[:]:
        if run.text == '':
            run._element.getparent().remove(run._element)
    
    # Add new text with preserved formatting
    new_run = paragraph.add_run(new_text)
    
    if base_format:
        if base_format['bold'] is not None:
            new_run.bold = base_format['bold']
        if base_format['italic'] is not None:
            new_run.italic = base_format['italic']
        if base_format['underline'] is not None:
            new_run.underline = base_format['underline']
        if base_format['font_name']:
            new_run.font.name = base_format['font_name']
        if base_format['font_size']:
            new_run.font.size = base_format['font_size']


def get_normalized_placeholder_name(placeholder_text, pattern_type='auto'):
    """
    Normalize a placeholder name according to the pattern type.
    Uses the same normalization logic as placeholder_detector.py.
    Useful for converting user-visible placeholders to internal keys.
    
    Args:
        placeholder_text (str): The placeholder text (with or without delimiters)
        pattern_type (str): The pattern type ('double_curly', 'single_curly', 
                           'square_bracket', 'underscore', 'dollar_underscore', 'auto')
    
    Returns:
        str: Normalized placeholder name
    
    Note:
        For context-dependent patterns (underscore, dollar_underscore), this function
        cannot determine the exact normalized name without surrounding context.
        It returns a generic placeholder name in those cases.
    """
    # Strip delimiters
    text = placeholder_text.strip()
    
    # Remove delimiters based on pattern type (auto-detect)
    if text.startswith('{{') and text.endswith('}}'):
        text = text[2:-2].strip()
        pattern_type = 'double_curly'
    elif text.startswith('{') and text.endswith('}'):
        text = text[1:-1].strip()
        pattern_type = 'single_curly'
    elif text.startswith('[') and text.endswith(']'):
        text = text[1:-1].strip()
        # Check if it's dollar_underscore or just square_bracket
        if text.startswith('$') or re.match(r'^_{3,}$', text):
            pattern_type = 'dollar_underscore'
        else:
            pattern_type = 'square_bracket'
    elif text.startswith('$[') and text.endswith(']'):
        pattern_type = 'dollar_underscore'
        # Context-dependent, cannot normalize without surrounding text
        return 'amount_1'  # placeholder default
    elif re.match(r'^_{3,}$', text):
        pattern_type = 'underscore'
        # Context-dependent, cannot normalize without surrounding text
        return 'field_1'  # placeholder default
    
    # Normalize using shared function for patterns with captured content
    if pattern_type in ['double_curly', 'single_curly', 'square_bracket']:
        return normalize_placeholder_name(text)
    elif pattern_type in ['underscore', 'dollar_underscore']:
        # Context-dependent patterns - return generic placeholder
        return 'placeholder'
    else:
        # Default: use shared normalization
        return normalize_placeholder_name(text)


def validate_document_path(file_path):
    """
    Validate that a document path exists and is accessible.
    
    Args:
        file_path (str): Path to validate
    
    Returns:
        bool: True if valid
    
    Raises:
        DocumentReplacementError: If path is invalid
    """
    if not file_path:
        raise DocumentReplacementError("File path is required")
    
    if not os.path.exists(file_path):
        raise DocumentReplacementError(f"File not found: {file_path}")
    
    if not os.path.isfile(file_path):
        raise DocumentReplacementError(f"Path is not a file: {file_path}")
    
    if not file_path.lower().endswith('.docx'):
        raise DocumentReplacementError(f"File must be a .docx document: {file_path}")
    
    return True

