"""
Unit tests for placeholder detection functionality
"""

import pytest
from lib.placeholder_detector import (
    detect_placeholders,
    detect_placeholders_with_context,
    reduce_false_positives,
    normalize_placeholder_name,
    get_placeholder_count,
    get_total_occurrences
)
from docx import Document
import tempfile
import os


@pytest.mark.unit
class TestPlaceholderDetection:
    """Test placeholder detection with various patterns."""
    
    def test_double_curly_braces(self):
        """Test detection of {{PLACEHOLDER}} pattern."""
        # Create a DOCX for this test
        doc = Document()
        doc.add_paragraph("This contract is between {{COMPANY_NAME}} and {{CLIENT_NAME}}.")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
        assert 'COMPANY_NAME' in result
        assert 'CLIENT_NAME' in result
        assert len(result) == 2
    
    def test_single_curly_braces(self):
        """Test detection of {PLACEHOLDER} pattern."""
        doc = Document()
        doc.add_paragraph("Payment due: {PAYMENT_DATE} Amount: {AMOUNT}")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
        assert 'PAYMENT_DATE' in result
        assert 'AMOUNT' in result
        assert len(result) == 2
    
    def test_square_brackets_uppercase(self):
        """Test detection of [PLACEHOLDER] pattern (uppercase)."""
        doc = Document()
        doc.add_paragraph("The agreement dated [CONTRACT_DATE] is hereby executed.")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
        assert 'CONTRACT_DATE' in result
        assert len(result) == 1
    
    def test_square_brackets_mixed_case(self):
        """Test detection of [Placeholder Name] pattern (mixed case)."""
        doc = Document()
        doc.add_paragraph("This [Date of Safe] shall be binding on [Party Name].")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
        assert 'Date_of_Safe' in result
        assert 'Party_Name' in result
        assert len(result) == 2
    
    def test_underscores(self):
        """Test detection of _____ pattern (5+ underscores)."""
        doc = Document()
        doc.add_paragraph("Signed by: _____ on date: _____")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
        assert 'UNDERSCORE_PLACEHOLDER' in result
        # Should group multiple underscore placeholders
        assert result['UNDERSCORE_PLACEHOLDER']['count'] == 2
    
    def test_dollar_brackets(self):
        """Test detection of $[_____] pattern."""
        doc = Document()
        doc.add_paragraph("Amount payable: $[_____] by $[_____]")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
        assert 'DOLLAR_BRACKET_PLACEHOLDER' in result
        assert result['DOLLAR_BRACKET_PLACEHOLDER']['count'] == 2
    
    def test_multiple_patterns_together(self):
        """Test detection of multiple pattern types in same text."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Contract between {{COMPANY}} and {CLIENT}.")
        doc.add_paragraph("Date: [CONTRACT_DATE]")
        doc.add_paragraph("Signature: _____")
        doc.add_paragraph("Amount: $[_____]")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
        assert 'COMPANY' in result
        assert 'CLIENT' in result
        assert 'CONTRACT_DATE' in result
        assert 'UNDERSCORE_PLACEHOLDER' in result
        assert 'DOLLAR_BRACKET_PLACEHOLDER' in result
        assert len(result) >= 5
    
    def test_repeated_placeholders(self):
        """Test that repeated placeholders are counted correctly."""
        doc = Document()
        doc.add_paragraph("{{NAME}} agrees that {{NAME}} will pay {{AMOUNT}}. {{AMOUNT}} is due on {{DATE}}.")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
        assert result['NAME']['count'] == 2
        assert result['AMOUNT']['count'] == 2
        assert result['DATE']['count'] == 1
    
    def test_case_insensitive_normalization(self):
        """Test that placeholder names are normalized (case-insensitive for curly braces)."""
        doc = Document()
        doc.add_paragraph("{{CompanyName}} and {{COMPANYNAME}} and {{company_name}}")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders(path)
        os.remove(path)
        
    def test_signature_line_without_underscores(self):
        """Address:/Email: with leader/tabs only should be detected in with_context."""
        doc = Document()
        doc.add_paragraph("Address:\t\t\t")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders_with_context(path)
        os.remove(path)
        placeholders = result['placeholders']
        assert 'address' in placeholders
        
    def test_bracket_underscore_removed(self):
        """[_____] should no longer be detected as a placeholder."""
        doc = Document()
        doc.add_paragraph("[_____]")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders_with_context(path)
        os.remove(path)
        assert 'blank' not in result['placeholders']
    
    def test_overlap_arbitration(self):
        """Higher-priority longer match should win; overlapping weaker matches dropped."""
        doc = Document()
        # Create overlapping patterns: [Company Name] contains potential underscore pattern
        doc.add_paragraph("[Company Name]")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders_with_context(path)
        os.remove(path)
        # Should detect square_bracket (priority 4), not underscore (priority 1)
        assert 'company_name' in result['placeholders']
        # Verify only one candidate for this span
        assert len(result['candidates']) == 1
        assert result['candidates'][0]['pattern_type'] == 'square_bracket'
    
    def test_context_and_instance_ids(self):
        """Ensure sentence context captured and stable instance IDs assigned."""
        doc = Document()
        doc.add_paragraph("This is the first sentence. The {{COMPANY}} will pay {{AMOUNT}}. This is the last sentence.")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders_with_context(path)
        os.remove(path)
        candidates = result['candidates']
        assert len(candidates) == 2
        # Check first candidate has context
        c1 = candidates[0]
        assert 'context' in c1
        assert 'prev' in c1['context']
        assert 'sentence' in c1['context']
        assert 'next' in c1['context']
        # Check instance IDs are stable and unique
        assert 'id' in c1
        assert c1['id'].startswith('p0-s')  # paragraph 0, start position
        assert candidates[0]['id'] != candidates[1]['id']
    
    def test_false_positive_filter_after_arbitration(self):
        """False-positive filter should still work after arbitration."""
        doc = Document()
        doc.add_paragraph("See [Section 2(a)] and {{COMPANY_NAME}}.")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            path = tmp.name
        result = detect_placeholders_with_context(path)
        os.remove(path)
        placeholders = result['placeholders']
        # Section reference should be filtered out
        assert 'section' not in placeholders
        assert 'section_2_a' not in placeholders
        # Real placeholder should remain
        assert 'company_name' in placeholders
        # Verify we have exactly 1 candidate (section ref was not added)
        assert len(result['candidates']) == 1
        assert result['candidates'][0]['normalized'] == 'company_name'


@pytest.mark.unit
class TestFalsePositiveFiltering:
    """Test false positive filtering for placeholder detection."""
    
    def test_filter_numeric_section_references(self):
        """Test filtering of numeric section references like [1], [2.3]."""
        placeholders = {
            'COMPANY_NAME': {'count': 1, 'positions': [(0, 12)]},
            '1': {'count': 1, 'positions': [(20, 21)]},
            '2.3': {'count': 1, 'positions': [(30, 33)]}
        }
        
        filtered = reduce_false_positives(placeholders, "")
        
        assert 'COMPANY_NAME' in filtered
        assert '1' not in filtered
        assert '2.3' not in filtered
    
    def test_filter_single_character(self):
        """Test filtering of single character placeholders."""
        placeholders = {
            'COMPANY_NAME': {'count': 1, 'positions': [(0, 12)]},
            'A': {'count': 1, 'positions': [(20, 21)]},
            'X': {'count': 1, 'positions': [(30, 31)]}
        }
        
        filtered = reduce_false_positives(placeholders, "")
        
        assert 'COMPANY_NAME' in filtered
        assert 'A' not in filtered
        assert 'X' not in filtered
    
    def test_keep_multi_word_placeholders(self):
        """Test that multi-word placeholders are kept."""
        placeholders = {
            'Company_Name': {'count': 1, 'positions': [(0, 12)]},
            'Date_of_Safe': {'count': 1, 'positions': [(20, 33)]},
            'A': {'count': 1, 'positions': [(40, 41)]}
        }
        
        filtered = reduce_false_positives(placeholders, "")
        
        assert 'Company_Name' in filtered
        assert 'Date_of_Safe' in filtered
        assert 'A' not in filtered
    
    def test_keep_short_but_frequent_placeholders(self):
        """Test that short placeholders appearing multiple times are kept."""
        placeholders = {
            'ID': {'count': 5, 'positions': [(0, 2), (10, 12), (20, 22), (30, 32), (40, 42)]},
            'X': {'count': 1, 'positions': [(50, 51)]}
        }
        
        filtered = reduce_false_positives(placeholders, "")
        
        # 'ID' appears 5 times, should be kept
        assert 'ID' in filtered
        # 'X' appears only once, should be filtered
        assert 'X' not in filtered


@pytest.mark.unit
class TestPlaceholderNormalization:
    """Test placeholder name normalization."""
    
    def test_normalize_curly_braces_uppercase(self):
        """Test normalization of curly brace placeholders to uppercase."""
        assert normalize_placeholder_name('Company Name') == 'COMPANY_NAME'
        assert normalize_placeholder_name('client-name') == 'CLIENT-NAME'
    
    def test_normalize_square_brackets_preserve_case(self):
        """Test normalization of square bracket placeholders preserves case."""
        assert normalize_placeholder_name('Date of Safe') == 'Date_of_Safe'
        assert normalize_placeholder_name('Party Name') == 'Party_Name'
    
    def test_normalize_spaces_to_underscores(self):
        """Test that spaces are converted to underscores."""
        assert normalize_placeholder_name('Company Name') == 'COMPANY_NAME'
        assert normalize_placeholder_name('Date of Safe') == 'Date_of_Safe'


@pytest.mark.unit
class TestUtilityFunctions:
    """Test utility functions for placeholder detection."""
    
    def test_get_placeholder_count(self):
        """Test counting unique placeholders."""
        placeholders = {
            'NAME': {'count': 2, 'positions': [(0, 4), (10, 14)]},
            'DATE': {'count': 1, 'positions': [(20, 24)]},
            'AMOUNT': {'count': 3, 'positions': [(30, 36), (40, 46), (50, 56)]}
        }
        
        count = get_placeholder_count(placeholders)
        assert count == 3
    
    def test_get_total_occurrences(self):
        """Test counting total placeholder occurrences."""
        placeholders = {
            'NAME': {'count': 2, 'positions': [(0, 4), (10, 14)]},
            'DATE': {'count': 1, 'positions': [(20, 24)]},
            'AMOUNT': {'count': 3, 'positions': [(30, 36), (40, 46), (50, 56)]}
        }
        
        total = get_total_occurrences(placeholders)
        assert total == 6  # 2 + 1 + 3
    
    def test_empty_placeholders(self):
        """Test handling of empty placeholder dictionary."""
        placeholders = {}
        
        assert get_placeholder_count(placeholders) == 0
        assert get_total_occurrences(placeholders) == 0

